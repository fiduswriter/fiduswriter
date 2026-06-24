"""
Management command to generate comprehensive test files for import testing.

Generates DOCX and ODT files covering all features supported by
Fidus Writer's import filters.
"""

import io
import os

from django.core.management.base import BaseCommand
from django.conf import settings


class Command(BaseCommand):
    help = "Generate comprehensive DOCX and ODT test files for import testing"

    def add_arguments(self, parser):
        parser.add_argument(
            "--output-dir",
            default=None,
            help="Output directory for test files (default: project tests fixtures dir)",
        )

    def handle(self, *args, **options):
        output_dir = options.get("output_dir")
        if output_dir is None:
            output_dir = os.path.join(
                settings.PROJECT_PATH,
                "..",
                "..",
                "fiduswriter-document",
                "test",
                "importer",
                "fixtures",
            )
        os.makedirs(output_dir, exist_ok=True)

        docx_path = os.path.join(output_dir, "comprehensive-test.docx")
        odt_path = os.path.join(output_dir, "comprehensive-test.odt")

        self.generate_docx(docx_path)
        self.stdout.write(self.style.SUCCESS(f"Generated DOCX: {docx_path}"))

        self.generate_odt(odt_path)
        self.stdout.write(self.style.SUCCESS(f"Generated ODT: {odt_path}"))

    # ------------------------------------------------------------------
    # DOCX generation (uses python-docx)
    # ------------------------------------------------------------------

    def generate_docx(self, path):
        from PIL import Image as PILImage
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import datetime

        doc = Document()

        # ---- METADATA ----
        doc.core_properties.title = (
            "Comprehensive Test Document for Fidus Writer Import"
        )
        doc.core_properties.author = "Jane Doe"
        doc.core_properties.keywords = (
            "testing, import, fiduswriter, comprehensive"
        )
        doc.core_properties.language = "en-US"

        # ---- H1 ----
        doc.add_heading("Test Document: Fidus Writer Import (DOCX)", level=1)

        # ---- Formatting ----
        doc.add_heading("Text Formatting", level=2)

        p = doc.add_paragraph()
        p.add_run("Bold: ").font.size = Pt(11)
        r = p.add_run("bold text")
        r.bold = True
        r.font.size = Pt(11)
        p.add_run(", Italic: ").font.size = Pt(11)
        r = p.add_run("italic text")
        r.italic = True
        r.font.size = Pt(11)
        p.add_run(", Underline: ").font.size = Pt(11)
        r = p.add_run("underlined")
        r.underline = True
        r.font.size = Pt(11)
        p.add_run(", Superscript: ").font.size = Pt(11)
        r = p.add_run("E=mc2")
        r.font.size = Pt(11)
        r2 = p.add_run("2")
        r2.font.superscript = True
        r2.font.size = Pt(11)
        p.add_run(", Subscript: ").font.size = Pt(11)
        r = p.add_run("H2O")
        r.font.size = Pt(11)
        r2 = p.add_run("2")
        r2.font.subscript = True
        r2.font.size = Pt(11)
        p.add_run(", Code: ").font.size = Pt(11)
        r = p.add_run("monospace")
        r.font.name = "Courier New"
        r.font.size = Pt(10)
        p.add_run(", Strikethrough: ").font.size = Pt(11)
        r = p.add_run("struck")
        r.font.strike = True
        r.font.size = Pt(11)

        # Alignments
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("Center-aligned paragraph").font.size = Pt(11)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run("Right-aligned paragraph").font.size = Pt(11)

        # ---- Lists ----
        doc.add_heading("Lists", level=2)
        doc.add_paragraph("Bullet item 1", style="List Bullet")
        doc.add_paragraph("Bullet item 2 with bold", style="List Bullet")
        doc.add_paragraph("Bullet item 3", style="List Bullet")

        doc.add_paragraph("Numbered item 1", style="List Number")
        doc.add_paragraph("Numbered item 2", style="List Number")
        doc.add_paragraph("Numbered item 3", style="List Number")

        # ---- Hyperlinks ----
        doc.add_heading("Hyperlinks", level=2)
        self._add_hyperlink(
            doc.add_paragraph(), "https://fiduswriter.org", "Fidus Writer"
        )

        # ---- Tables ----
        doc.add_heading("Tables", level=2)
        table = doc.add_table(rows=3, cols=3)
        table.style = "Table Grid"
        for i, h in enumerate(["Header 1", "Header 2", "Header 3"]):
            cell = table.rows[0].cells[i]
            cell.text = h
            for par in cell.paragraphs:
                for run in par.runs:
                    run.bold = True
        for row_idx in range(1, 3):
            for col_idx in range(3):
                table.rows[row_idx].cells[
                    col_idx
                ].text = f"Row {row_idx}, C{col_idx+1}"

        # Table with merged cells
        doc.add_paragraph()
        table2 = doc.add_table(rows=3, cols=3)
        table2.style = "Table Grid"
        for i, h in enumerate(["Name", "Value", "Notes"]):
            table2.rows[0].cells[i].text = h
        table2.rows[1].cells[0].text = "Item 1"
        table2.rows[1].cells[1].text = "42"
        table2.rows[1].cells[2].text = "First"
        table2.rows[2].cells[0].merge(table2.rows[2].cells[2])
        table2.rows[2].cells[0].text = "Merged cell spanning all columns"

        # ---- Footnotes ----
        doc.add_heading("Footnotes", level=2)
        p = doc.add_paragraph()
        p.add_run("Paragraph with").font.size = Pt(11)
        self._add_footnote(doc, p, "First footnote content.")
        p.add_run(" a footnote.").font.size = Pt(11)

        # ---- Images ----
        doc.add_heading("Images", level=2)
        doc.add_paragraph().add_run("PNG image below:").font.size = Pt(11)
        img = PILImage.new("RGB", (100, 50), (100, 150, 200))
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)
        doc.add_picture(buf, width=Inches(2), height=Inches(1))

        # ---- Tracked Changes ----
        doc.add_heading("Tracked Changes", level=2)
        p = doc.add_paragraph()
        self._add_tracked_insertion(
            p, "Inserted tracked text. ", "Jane Doe", datetime.datetime.now()
        )
        p.add_run("Original stable text.").font.size = Pt(11)
        self._add_tracked_deletion(
            p, "Deleted tracked text.", "Jane Doe", datetime.datetime.now()
        )

        # ---- Block Quote ----
        doc.add_heading("Block Quotes", level=2)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.right_indent = Inches(0.5)
        p.add_run("This is a block quote for testing.").italic = True

        # ---- Comments ----
        doc.add_heading("Comments", level=2)
        p = doc.add_paragraph()
        p.add_run("Paragraph with ").font.size = Pt(11)
        self._add_comment(doc, p, "commented", "Test comment.", "Jane Doe")
        p.add_run(" word.").font.size = Pt(11)

        # ---- Citations ----
        doc.add_heading("Citations", level=2)
        p = doc.add_paragraph()
        p.add_run("This paragraph contains a citation: ").font.size = Pt(11)
        # Add a citation field code (Zotero-style ADDIN citation)
        self._add_citation_field(
            doc,
            p,
            "Doe, 2012",
            "Doe2012",
            "Doe, John. (2012). My title. In My publication title.",
        )

        doc.save(path)

    @staticmethod
    def _add_hyperlink(paragraph, url, text):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hl = OxmlElement("w:hyperlink")
        hl.set(qn("r:id"), r_id)
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "Hyperlink")
        rPr.append(rStyle)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        r.append(t)
        hl.append(r)
        paragraph._p.append(hl)

    @staticmethod
    def _add_footnote(doc, paragraph, text):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        # Get footnotes part
        footnotes_part = None
        for rel in doc.part.rels.values():
            if "footnotes" in rel.reltype:
                footnotes_part = rel.target_part
                break
        if footnotes_part is None:
            return
        existing_ids = [
            int(fn.get(qn("w:id")))
            for fn in footnotes_part._element.findall(qn("w:footnote"))
            if fn.get(qn("w:id")) not in ("0", "-1", None)
        ]
        fn_id = max(existing_ids + [1]) + 1
        # Reference in body
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "FootnoteReference")
        rPr.append(rStyle)
        r.append(rPr)
        ref = OxmlElement("w:footnoteReference")
        ref.set(qn("w:id"), str(fn_id))
        r.append(ref)
        paragraph._p.append(r)
        # Footnote content
        fn = OxmlElement("w:footnote")
        fn.set(qn("w:id"), str(fn_id))
        fn_p = OxmlElement("w:p")
        fn_pPr = OxmlElement("w:pPr")
        fn_pStyle = OxmlElement("w:pStyle")
        fn_pStyle.set(qn("w:val"), "FootnoteText")
        fn_pPr.append(fn_pStyle)
        fn_p.append(fn_pPr)
        fn_r = OxmlElement("w:r")
        fn_t = OxmlElement("w:t")
        fn_t.text = f" {text}"
        fn_r.append(fn_t)
        fn_p.append(fn_r)
        fn.append(fn_p)
        footnotes_part._element.append(fn)

    @staticmethod
    def _add_tracked_insertion(paragraph, text, author, date):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        ins = OxmlElement("w:ins")
        ins.set(qn("w:id"), str(abs(hash(text + "ins")) % 100000))
        ins.set(qn("w:author"), author)
        ins.set(qn("w:date"), date.isoformat() + "Z")
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = text
        r.append(t)
        ins.append(r)
        paragraph._p.append(ins)

    @staticmethod
    def _add_tracked_deletion(paragraph, text, author, date):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        dele = OxmlElement("w:del")
        dele.set(qn("w:id"), str(abs(hash(text + "del")) % 100000 + 50000))
        dele.set(qn("w:author"), author)
        dele.set(qn("w:date"), date.isoformat() + "Z")
        r = OxmlElement("w:r")
        dt = OxmlElement("w:delText")
        dt.text = text
        r.append(dt)
        dele.append(r)
        paragraph._p.append(dele)

    @staticmethod
    def _add_comment(doc, paragraph, marked_text, comment_text, author):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        import datetime

        comments_part = None
        for rel in doc.part.rels.values():
            if "comments" in rel.reltype:
                comments_part = rel.target_part
                break
        if comments_part is None:
            return
        existing_ids = [
            int(c.get(qn("w:id")))
            for c in comments_part._element.findall(qn("w:comment"))
            if c.get(qn("w:id"))
        ]
        cid = max(existing_ids + [0]) + 1
        # Range start
        rs = OxmlElement("w:commentRangeStart")
        rs.set(qn("w:id"), str(cid))
        paragraph._p.append(rs)
        # Text
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = marked_text
        r.append(t)
        paragraph._p.append(r)
        # Range end
        re = OxmlElement("w:commentRangeEnd")
        re.set(qn("w:id"), str(cid))
        paragraph._p.append(re)
        # Reference
        cr = OxmlElement("w:r")
        crPr = OxmlElement("w:rPr")
        cs = OxmlElement("w:rStyle")
        cs.set(qn("w:val"), "CommentReference")
        crPr.append(cs)
        cr.append(crPr)
        cref = OxmlElement("w:commentReference")
        cref.set(qn("w:id"), str(cid))
        cr.append(cref)
        paragraph._p.append(cr)
        # Comment content
        ce = OxmlElement("w:comment")
        ce.set(qn("w:id"), str(cid))
        ce.set(qn("w:author"), author)
        ce.set(qn("w:date"), datetime.datetime.now().isoformat() + "Z")
        ce.set(qn("w:initials"), "JD")
        cp = OxmlElement("w:p")
        cpr = OxmlElement("w:r")
        ct = OxmlElement("w:t")
        ct.text = comment_text
        cpr.append(ct)
        cp.append(cpr)
        ce.append(cp)
        comments_part._element.append(ce)

    def _add_citation_field(
        self, doc, paragraph, display_text, citation_id, bib_entry
    ):
        """Add a citation field to a DOCX paragraph.

        This creates a field code that mimics Zotero/Mendeley style citations
        so that biblatex-csl-converter can parse it.
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        # Create the field code for a citation
        # Format: ADDIN ZOTERO_ITEM CSL_CITATION{json}
        # Note: No space between CSL_CITATION and the JSON object
        json_part = (
            '{"citationID": "' + citation_id + '", '
            '"properties": {"formattedCitation": "' + display_text + '", '
            '"plainCitation": "' + display_text + '", "noteIndex": 0}, '
            '"citationItems": [{"id": "' + citation_id + '", '
            '"uris": ["http://zotero.org/items/' + citation_id + '"], '
            '"itemData": {"title": "My title", '
            '"author": [{"family": "Doe", "given": "John"}], '
            '"date": "2012"}}]'
        )
        field_code = "ADDIN ZOTERO_ITEM CSL_CITATION" + json_part

        # Create the field structure
        p = paragraph._p

        # Field begin
        fldChar_begin = OxmlElement("w:fldChar")
        fldChar_begin.set(qn("w:fldCharType"), "begin")
        p.append(fldChar_begin)

        # Field instruction text
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = field_code
        p.append(instrText)

        # Field separate
        fldChar_sep = OxmlElement("w:fldChar")
        fldChar_sep.set(qn("w:fldCharType"), "separate")
        p.append(fldChar_sep)

        # Display text (citation text)
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = display_text
        r.append(t)
        p.append(r)

        # Field end
        fldChar_end = OxmlElement("w:fldChar")
        fldChar_end.set(qn("w:fldCharType"), "end")
        p.append(fldChar_end)

    # ------------------------------------------------------------------
    # ODT generation (manual XML via zipfile)
    # ------------------------------------------------------------------

    def generate_odt(self, path):
        """Create a minimal valid ODT file with raw XML."""
        import zipfile

        content_xml = """<?xml version="1.0" encoding="UTF-8"?>
<office:document-content
  xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0"
  xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0"
  xmlns:style="urn:oasis:names:tc:opendocument:xmlns:style:1.0"
  xmlns:fo="urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0"
  xmlns:dc="http://purl.org/dc/elements/1.1/"
  xmlns:meta="urn:oasis:names:tc:opendocument:xmlns:meta:1.0"
  xmlns:table="urn:oasis:names:tc:opendocument:xmlns:table:1.0"
  xmlns:draw="urn:oasis:names:tc:opendocument:xmlns:drawing:1.0"
  xmlns:xlink="http://www.w3.org/1999/xlink"
  xmlns:svg="urn:oasis:names:tc:opendocument:xmlns:svg-compatible:1.0"
  office:version="1.2">
<office:scripts/>
<office:font-face-decls>
  <style:font-face style:name="Courier New" svg:font-family="Courier New"/>
</office:font-face-decls>
<office:automatic-styles>
  <style:style style:name="T1" style:family="text">
    <style:text-properties fo:font-weight="bold"/>
  </style:style>
  <style:style style:name="T2" style:family="text">
    <style:text-properties fo:font-style="italic"/>
  </style:style>
  <style:style style:name="T3" style:family="text">
    <style:text-properties style:text-underline-type="single"/>
  </style:style>
  <style:style style:name="T4" style:family="text">
    <style:text-properties fo:font-family="Courier New" fo:font-size="10pt"/>
  </style:style>
  <style:style style:name="T5" style:family="text">
    <style:text-properties style:text-position="super 58%"/>
  </style:style>
  <style:style style:name="T6" style:family="text">
    <style:text-properties style:text-position="sub 58%"/>
  </style:style>
  <style:style style:name="P1" style:family="paragraph" style:parent-style-name="Standard">
    <style:paragraph-properties fo:text-align="center"/>
  </style:style>
  <style:style style:name="P2" style:family="paragraph" style:parent-style-name="Standard">
    <style:paragraph-properties fo:text-align="end"/>
  </style:style>
</office:automatic-styles>
<office:body>
<office:text>
  <text:h text:outline-level="1">Test Document: Fidus Writer Import (ODT)</text:h>

  <text:h text:outline-level="2">Text Formatting</text:h>
  <text:p>Normal paragraph with
    <text:span text:style-name="T1">bold</text:span>,
    <text:span text:style-name="T2">italic</text:span>,
    <text:span text:style-name="T3">underlined</text:span>,
    <text:span text:style-name="T4">monospace code</text:span>,
    <text:span text:style-name="T5">superscript</text:span>, and
    <text:span text:style-name="T6">subscript</text:span> text.
  </text:p>
  <text:p text:style-name="P1">Center-aligned paragraph</text:p>
  <text:p text:style-name="P2">Right-aligned paragraph</text:p>

  <text:h text:outline-level="2">Lists</text:h>
  <text:list>
    <text:list-item><text:p>Bullet item 1</text:p></text:list-item>
    <text:list-item><text:p>Bullet item 2</text:p></text:list-item>
    <text:list-item><text:p>Bullet item 3</text:p></text:list-item>
  </text:list>
  <text:list text:style-name="L1">
    <text:list-item><text:p>Numbered item 1</text:p></text:list-item>
    <text:list-item><text:p>Numbered item 2</text:p></text:list-item>
    <text:list-item><text:p>Numbered item 3</text:p></text:list-item>
  </text:list>

  <text:h text:outline-level="2">Hyperlinks</text:h>
  <text:p>Visit <text:a xlink:type="simple" xlink:href="https://fiduswriter.org">Fidus Writer</text:a></text:p>

  <text:h text:outline-level="2">Tables</text:h>
  <table:table table:name="Table1">
    <table:table-column table:number-columns-repeated="3"/>
    <table:table-row>
      <table:table-cell office:value-type="string"><text:p>H1</text:p></table:table-cell>
      <table:table-cell office:value-type="string"><text:p>H2</text:p></table:table-cell>
      <table:table-cell office:value-type="string"><text:p>H3</text:p></table:table-cell>
    </table:table-row>
    <table:table-row>
      <table:table-cell office:value-type="string"><text:p>A</text:p></table:table-cell>
      <table:table-cell office:value-type="string"><text:p>B</text:p></table:table-cell>
      <table:table-cell office:value-type="string"><text:p>C</text:p></table:table-cell>
    </table:table-row>
    <table:table-row>
      <table:table-cell table:number-columns-spanned="3" office:value-type="string">
        <text:p>Merged cell</text:p>
      </table:table-cell>
    </table:table-row>
  </table:table>

  <text:h text:outline-level="2">Footnotes</text:h>
  <text:p>Paragraph with a
    <text:note text:id="ftn1" text:note-class="footnote">
      <text:note-citation>1</text:note-citation>
      <text:note-body><text:p>ODT footnote content.</text:p></text:note-body>
    </text:note>
    footnote.
  </text:p>

  <text:h text:outline-level="2">Block Quote</text:h>
  <text:p>This is a quoted paragraph for testing.</text:p>

  <text:h text:outline-level="2">Tracked Changes</text:h>
  <text:p>Original text with <text:span text:style-name="T1">formatting</text:span>.</text:p>

  <text:h text:outline-level="2">Comments</text:h>
  <text:p>This paragraph has a commented word in it.</text:p>

  <text:h text:outline-level="2">Citations</text:h>
  <text:p>Paragraph with a <text:reference-mark-start text:name="ZOTERO_ITEM CSL_CITATION {&quot;citationID&quot;: &quot;Doe2012&quot;, &quot;properties&quot;: {&quot;formattedCitation&quot;: &quot;(Doe, 2012)&quot;, &quot;plainCitation&quot;: &quot;(Doe, 2012)&quot;, &quot;noteIndex&quot;: 0}, &quot;citationItems&quot;: [{&quot;id&quot;: &quot;Doe2012&quot;, &quot;uris&quot;: [&quot;http://zotero.org/items/Doe2012&quot;], &quot;itemData&quot;: {&quot;title&quot;: &quot;My title&quot;, &quot;author&quot;: [{&quot;family&quot;: &quot;Doe&quot;, &quot;given&quot;: &quot;John&quot;}], &quot;date&quot;: &quot;2012&quot;}}] RNDjURflxg9F1"/>citation (Doe, 2012)<text:reference-mark-end text:name="ZOTERO_ITEM CSL_CITATION {&quot;citationID&quot;: &quot;Doe2012&quot;, &quot;properties&quot;: {&quot;formattedCitation&quot;: &quot;(Doe, 2012)&quot;, &quot;plainCitation&quot;: &quot;(Doe, 2012)&quot;, &quot;noteIndex&quot;: 0}, &quot;citationItems&quot;: [{&quot;id&quot;: &quot;Doe2012&quot;, &quot;uris&quot;: [&quot;http://zotero.org/items/Doe2012&quot;], &quot;itemData&quot;: {&quot;title&quot;: &quot;My title&quot;, &quot;author&quot;: [{&quot;family&quot;: &quot;Doe&quot;, &quot;given&quot;: &quot;John&quot;}], &quot;date&quot;: &quot;2012&quot;}}] RNDjURflxg9F1"/> in the text.</text:p>

</office:text>
</office:body>
</office:document-content>"""

        styles_xml = """<?xml version="1.0" encoding="UTF-8"?>
<office:document-styles
  xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0"
  xmlns:style="urn:oasis:names:tc:opendocument:xmlns:style:1.0"
  xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0"
  xmlns:fo="urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0"
  office:version="1.2">
<office:styles>
  <style:style style:name="Standard" style:family="paragraph">
    <style:text-properties fo:font-size="11pt"/>
  </style:style>
  <style:style style:name="Heading" style:family="paragraph">
    <style:text-properties fo:font-size="14pt" fo:font-weight="bold"/>
  </style:style>
</office:styles>
<office:automatic-styles>
  <text:list-style style:name="L1">
    <text:list-level-style-number text:level="1" style:num-format="1">
      <style:list-level-properties text:list-level-position-and-space-mode="label-alignment">
        <style:list-level-label-alignment text:label-followed-by="listtab"/>
      </style:list-level-properties>
    </text:list-level-style-number>
  </text:list-style>
</office:automatic-styles>
</office:document-styles>"""

        meta_xml = """<?xml version="1.0" encoding="UTF-8"?>
<office:document-meta
  xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0"
  xmlns:dc="http://purl.org/dc/elements/1.1/"
  xmlns:meta="urn:oasis:names:tc:opendocument:xmlns:meta:1.0"
  office:version="1.2">
<office:meta>
  <dc:title>Comprehensive Test Document for Fidus Writer Import (ODT)</dc:title>
  <dc:creator>Jane Doe</dc:creator>
  <dc:subject>testing, import, fiduswriter, odt</dc:subject>
  <dc:language>en-US</dc:language>
  <meta:creation-date>2024-01-01T00:00:00</meta:creation-date>
</office:meta>
</office:document-meta>"""

        manifest_xml = """<?xml version="1.0" encoding="UTF-8"?>
<manifest:manifest
  xmlns:manifest="urn:oasis:names:tc:opendocument:xmlns:manifest:1.0"
  manifest:version="1.2">
  <manifest:file-entry manifest:full-path="/" manifest:version="1.2"
    manifest:media-type="application/vnd.oasis.opendocument.text"/>
  <manifest:file-entry manifest:full-path="content.xml"
    manifest:media-type="text/xml"/>
  <manifest:file-entry manifest:full-path="styles.xml"
    manifest:media-type="text/xml"/>
  <manifest:file-entry manifest:full-path="meta.xml"
    manifest:media-type="text/xml"/>
</manifest:manifest>"""

        with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("mimetype", "application/vnd.oasis.opendocument.text")
            zf.writestr("content.xml", content_xml)
            zf.writestr("styles.xml", styles_xml)
            zf.writestr("meta.xml", meta_xml)
            zf.writestr("META-INF/manifest.xml", manifest_xml)
